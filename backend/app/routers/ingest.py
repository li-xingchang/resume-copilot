"""
POST /ingest — parse a PDF/DOCX resume into career_facts with embeddings.

Flow:
  1. Verify Clerk JWT; upsert user row on first call
  2. Extract raw text from uploaded file (pdfplumber for PDF, python-docx for DOCX)
  3. LLM parses text into atomic career facts
  4. Embed each fact in batch (one API call to text-embedding-3-small)
  5. Persist to career_facts; return for user verification on /onboard
"""
import io
import uuid

import pdfplumber
from docx import Document
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from sqlalchemy.dialects.postgresql import insert as pg_insert
from sqlalchemy.ext.asyncio import AsyncSession

from app.auth import get_verified_user_id
from app.database import get_db
from app.models.tables import CareerFact, User
from app.schemas.api import CareerFactOut, IngestResponse
from app.services import embeddings as emb_svc
from app.services import llm as llm_svc

router = APIRouter()


def _extract_text_pdf(data: bytes) -> str:
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def _extract_text_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


async def _upsert_user(db: AsyncSession, clerk_id: str, email: str = "") -> uuid.UUID:
    """
    Create user row on first ingest if it doesn't exist yet.
    Uses clerk_id as the stable identity; UUID is generated deterministically
    from the clerk_id so it's consistent across upserts.
    """
    # Deterministic UUID from clerk_id keeps the PK stable across retries
    user_id = uuid.uuid5(uuid.NAMESPACE_URL, f"clerk:{clerk_id}")
    await db.execute(
        pg_insert(User)
        .values(id=user_id, clerk_id=clerk_id, email=email or f"{clerk_id}@clerk.local")
        .on_conflict_do_nothing(index_elements=["clerk_id"])
    )
    await db.flush()
    return user_id


@router.post("/ingest", response_model=IngestResponse)
async def ingest(
    file: UploadFile = File(...),
    email: str = Form(""),                          # optional; Clerk passes this
    db: AsyncSession = Depends(get_db),
    clerk_id: str = Depends(get_verified_user_id),  # JWT-verified
) -> IngestResponse:
    user_id = await _upsert_user(db, clerk_id, email)

    data = await file.read()
    content_type = file.content_type or ""

    if "pdf" in content_type or (file.filename or "").endswith(".pdf"):
        raw_text = _extract_text_pdf(data)
    elif "word" in content_type or (file.filename or "").endswith(".docx"):
        raw_text = _extract_text_docx(data)
    else:
        raise HTTPException(status_code=415, detail="Only PDF and DOCX are supported")

    if not raw_text.strip():
        raise HTTPException(status_code=422, detail="Could not extract text from file")

    raw_facts = await llm_svc.parse_resume_to_facts(raw_text)
    if not raw_facts:
        raise HTTPException(status_code=422, detail="No career facts extracted from resume")

    texts = [f["canonical_text"] for f in raw_facts]
    vectors = await emb_svc.embed_batch(texts)

    saved: list[CareerFact] = []
    for raw, vec in zip(raw_facts, vectors):
        fact = CareerFact(
            user_id=user_id,
            type=raw.get("type", "achievement"),
            canonical_text=raw["canonical_text"],
            metric_json=raw.get("metric_json"),
            embedding=vec,
            is_verified=False,
            source_section=raw.get("source_section"),
        )
        db.add(fact)
        saved.append(fact)

    await db.flush()

    return IngestResponse(
        fact_count=len(saved),
        facts=[
            CareerFactOut(
                id=f.id,
                type=f.type,
                canonical_text=f.canonical_text,
                metric_json=f.metric_json,
                is_verified=f.is_verified,
                source_section=f.source_section,
            )
            for f in saved
        ],
    )
